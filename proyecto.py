import mysql.connector
import argparse
import re
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import os

class SistemaVentas:
    def __init__(self):
      
        self.conexion = mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="algoritmos"
        )
        self.cursor = self.conexion.cursor()

    def cerrar_conexion(self):
        self.cursor.close()
        self.conexion.close()

    # Control de Inventario
    def listar_productos(self):
        self.cursor.execute("SELECT codigo, nombre, existencia, proveedor, precio FROM inventario")
        for (codigo, nombre, existencia, proveedor, precio) in self.cursor:
            print(f"Código: {codigo}, Nombre: {nombre}, Existencia: {existencia}, Proveedor: {proveedor}, Precio: {precio}")

    def crear_producto(self, codigo, nombre, existencia, proveedor, precio):
        query = "INSERT INTO inventario (codigo, nombre, existencia, proveedor, precio) VALUES (%s, %s, %s, %s, %s)"
        values = (codigo, nombre, existencia, proveedor, precio)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Producto '{nombre}' creado con éxito.")

    def actualizar_producto(self, codigo, nombre, existencia, proveedor, precio):
        query = "UPDATE inventario SET nombre = %s, existencia = %s, proveedor = %s, precio = %s WHERE codigo = %s"
        values = (nombre, existencia, proveedor, precio, codigo)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Producto '{nombre}' actualizado con éxito.")

    def editar_existencias(self, codigo, cantidad):
        query = "UPDATE inventario SET existencia = existencia + %s WHERE codigo = %s"
        values = (cantidad, codigo)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Existencias del producto editadas con éxito.")

    def eliminar_producto(self, codigo):
        query = "DELETE FROM inventario WHERE codigo = %s"
        self.cursor.execute(query, (codigo,))
        self.conexion.commit()
        print(f"Producto eliminado con éxito.")

    # Control de Clientes
    def listar_clientes(self):
        self.cursor.execute("SELECT codigo, nombre, direccion FROM clientes")
        for (codigo, nombre, direccion) in self.cursor:
            print(f"Código: {codigo}, Nombre: {nombre}, Dirección: {direccion}")

    def crear_cliente(self, codigo, nombre, direccion):
        query = "INSERT INTO clientes (codigo, nombre, direccion) VALUES (%s, %s, %s)"
        values = (codigo, nombre, direccion)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Cliente '{nombre}' creado con éxito.")

    def editar_cliente(self, codigo, nombre, direccion):
        query = "UPDATE clientes SET nombre = %s, direccion = %s WHERE codigo = %s"
        values = (nombre, direccion, codigo)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Cliente '{nombre}' actualizado con éxito.")

    def eliminar_cliente(self, codigo):
        query = "DELETE FROM clientes WHERE codigo = %s"
        self.cursor.execute(query, (codigo,))
        self.conexion.commit()
        print(f"Cliente eliminado con éxito.")

    # Control de Ventas
    def listar_ventas(self):
        self.cursor.execute("SELECT codigo_producto, codigo_cliente, cantidad, total FROM ventas")
        for (codigo_producto, codigo_cliente, cantidad, total) in self.cursor:
            print(f"Código de producto: {codigo_producto}, Código de cliente: {codigo_cliente}, Cantidad: {cantidad}, Total: {total}")

    def crear_venta(self, codigo_producto, codigo_cliente, cantidad, total):
        query = "INSERT INTO ventas (codigo_producto, codigo_cliente, cantidad, total) VALUES (%s, %s, %s, %s)"
        values = (codigo_producto, codigo_cliente, cantidad, total)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print("Venta registrada con éxito.")

    def anular_venta(self, codigo_producto, codigo_cliente):
        query = "DELETE FROM ventas WHERE codigo_producto = %s AND codigo_cliente = %s"
        values = (codigo_producto, codigo_cliente)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print("Venta anulada con éxito.")

    # Reportes Básicos
    def ventas_por_cliente(self, codigo_cliente):
        query = "SELECT codigo_producto, cantidad, total FROM ventas WHERE codigo_cliente = %s"
        self.cursor.execute(query, (codigo_cliente,))
        print(f"Ventas del cliente (Código {codigo_cliente}):")
        for (codigo_producto, cantidad, total) in self.cursor:
            print(f"Código de producto: {codigo_producto}, Cantidad: {cantidad}, Total: {total}")

    def ventas_por_producto(self, codigo_producto):
        query = "SELECT codigo_cliente, cantidad, total FROM ventas WHERE codigo_producto = %s"
        self.cursor.execute(query, (codigo_producto,))
        print(f"Ventas del producto (Código {codigo_producto}):")
        for (codigo_cliente, cantidad, total) in self.cursor:
            print(f"Código de cliente: {codigo_cliente}, Cantidad: {cantidad}, Total: {total}")

    def obtener_ventas_por_cliente(self, codigo_cliente):
        query = "SELECT codigo_producto, codigo_cliente, cantidad, total FROM ventas WHERE codigo_cliente = %s"
        self.cursor.execute(query, (codigo_cliente,))
        return self.cursor.fetchall()

    def obtener_ventas_por_producto(self, codigo_producto):
        query = "SELECT codigo_producto, codigo_cliente, cantidad, total FROM ventas WHERE codigo_producto = %s"
        self.cursor.execute(query, (codigo_producto,))
        return self.cursor.fetchall()

    def generar_reporte_ventas(self, ventas, nombre_reporte):
        doc = Document()
        doc.add_heading(f'Reporte de Ventas - ', 0)

        for venta in ventas:
            doc.add_paragraph(f'Producto: {venta[0]}')
            doc.add_paragraph(f'Cliente: {venta[1]}')
            doc.add_paragraph(f'Cantidad: {venta[2]}')
            doc.add_paragraph(f'Total: {venta[3]}')
            doc.add_paragraph('')

        doc.save(f'reporte_ventas.docx')
        print(f'Reporte generado y guardado como reporte_ventas.docx')


    def enviar_correo_con_adjunto(self, destinatario, asunto, mensaje, archivo_adjunto):
      
        servidor_smtp = "smtp.gmail.com" 
        puerto_smtp = 587  
        correo_emisor = "leoestuarlem@gmail.com"  
        clave_emisor = ""  #Se elimino la contraseña por seguridad

        
        servidor = smtplib.SMTP(servidor_smtp, puerto_smtp)

      
        servidor.starttls()
        servidor.login(correo_emisor, clave_emisor)

       
        mensaje_correo = MIMEMultipart()
        mensaje_correo["From"] = correo_emisor
        mensaje_correo["To"] = destinatario
        mensaje_correo["Subject"] = asunto

       
        with open(archivo_adjunto, "rb") as adjunto:
            part = MIMEApplication(adjunto.read(), Name=os.path.basename(archivo_adjunto))
            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(archivo_adjunto)}"'
            mensaje_correo.attach(part)

      
        mensaje_correo.attach(MIMEText(mensaje, "plain"))

       
        servidor.sendmail(correo_emisor, destinatario, mensaje_correo.as_string())

        servidor.quit()
    

def mostrar_menu():
    print("Menu:")
    print("1. Control de Inventario")
    print("2. Control de Clientes")
    print("3. Control de Ventas")
    print("4. Reportes Básicos")
    print("5. Salir")

def main():
    parser = argparse.ArgumentParser(description="Sistema de Ventas")
    parser.add_argument("--ayuda", help="Mostrar ayuda", action="store_true")
    parser.add_argument("--inventario", help="Control de Inventario", action="store_true")
    parser.add_argument("--listar", help="Listar productos", action="store_true")
    parser.add_argument("--crear", help="Crear producto")
    parser.add_argument("--actualizar", help="Actualizar producto")
    parser.add_argument("--existencia", help="Editar existencias de producto")
    parser.add_argument("--eliminar", help="Eliminar producto")

    args = parser.parse_args()
    sistema = SistemaVentas()

    if args.ayuda:
        parser.print_help()
    elif args.inventario:
        if args.listar:
            sistema.listar_productos()
        elif args.crear:
            if len(args.crear) == 5:
                codigo, nombre, existencia, proveedor, precio = args.crear
                sistema.crear_producto(int(codigo), nombre, int(existencia), proveedor, float(precio))
            else:
                print("Error: se requieren 5 argumentos para crear un producto.")
        elif args.actualizar:
            codigo, nombre, existencia, proveedor, precio = args.actualizar.split()
            sistema.actualizar_producto(codigo, nombre, int(existencia), proveedor, float(precio))
        elif args.existencia:
            codigo, cantidad = args.existencia.split()
            sistema.editar_existencias(codigo, int(cantidad))
        elif args.eliminar:
            codigo = args.eliminar
            sistema.eliminar_producto(codigo)

            

     
 

    sistema = SistemaVentas()

    
    while True:
        mostrar_menu()
        opcion = input("Elija una opción: ")

        if opcion == "1":
            # Control de Inventario
            print("Control de Inventario:")
            print("a. Listar productos")
            print("b. Crear producto")
            print("c. Actualizar producto")
            print("d. Editar existencias de producto")
            print("e. Eliminar producto")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                sistema.listar_productos()
            elif sub_opcion == "b":
                codigo = input("Código: ")
                nombre = input("Nombre: ")
                existencia = int(input("Existencia: "))
                proveedor = input("Proveedor: ")
                precio = float(input("Precio: "))
                sistema.crear_producto(codigo, nombre, existencia, proveedor, precio)
            elif sub_opcion == "c":
                codigo = input("Código del producto a actualizar: ")
                nombre = input("Nuevo nombre: ")
                existencia = int(input("Nueva existencia: "))
                proveedor = input("Nuevo proveedor: ")
                precio = float(input("Nuevo precio: "))
                sistema.actualizar_producto(codigo, nombre, existencia, proveedor, precio)
            elif sub_opcion == "d":
                codigo = input("Código del producto a editar existencias: ")
                cantidad = int(input("Cantidad a agregar/reducir: "))
                sistema.editar_existencias(codigo, cantidad)
            elif sub_opcion == "e":
                codigo = input("Código del producto a eliminar: ")
                sistema.eliminar_producto(codigo)

        elif opcion == "2":
            # Control de Clientes
            print("Control de Clientes:")
            print("a. Listar clientes")
            print("b. Crear cliente")
            print("c. Editar cliente")
            print("d. Eliminar cliente")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                sistema.listar_clientes()
            elif sub_opcion == "b":
                codigo = input("Código: ")
                nombre = input("Nombre: ")
                direccion = input("Dirección: ")
                sistema.crear_cliente(codigo, nombre, direccion)
            elif sub_opcion == "c":
                codigo = input("Código del cliente a editar: ")
                nombre = input("Nuevo nombre: ")
                direccion = input("Nueva dirección: ")
                sistema.editar_cliente(codigo, nombre, direccion)
            elif sub_opcion == "d":
                codigo = input("Código del cliente a eliminar: ")
                sistema.eliminar_cliente(codigo)


        elif opcion == "3":
            # Control de Ventas
            print("Control de Ventas:")
            print("a. Listar ventas")
            print("b. Crear venta")
            print("c. Anular venta")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                sistema.listar_ventas()
            elif sub_opcion == "b":
                codigo_producto = input("Código de producto: ")
                codigo_cliente = input("Código de cliente: ")
                cantidad = int(input("Cantidad de productos: "))
                total = float(input("Total de venta: "))
                sistema.crear_venta(codigo_producto, codigo_cliente, cantidad, total)
            elif sub_opcion == "c":
                codigo_producto = input("Código del producto de la venta a anular: ")
                codigo_cliente = input("Código del cliente de la venta a anular: ")
                sistema.anular_venta(codigo_producto, codigo_cliente)
        elif opcion == "4":
            # Reportes Básicos
            print("Reportes Básicos:")
            print("a. Ventas por cliente")
            print("b. Ventas por producto")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                codigo_cliente = input("Código del cliente: ")
                sistema.ventas_por_cliente(codigo_cliente)
                ventas = sistema.obtener_ventas_por_cliente(codigo_cliente)
                sistema.generar_reporte_ventas(ventas, f'reporte_ventas')
                archivo_adjunto = f'reporte_ventas.docx'
                destinatario = '' #Se elimino el correo destinatario por seguridad 
                asunto = 'Reporte de Ventas por Cliente'
                mensaje = 'Adjunto se encuentra el reporte de ventas por cliente.'
                sistema.enviar_correo_con_adjunto(destinatario, asunto, mensaje, archivo_adjunto)
            elif sub_opcion == "b":
                codigo_producto = input("Código del producto: ")
                sistema.ventas_por_producto(codigo_producto)
                ventas = sistema.obtener_ventas_por_producto(codigo_producto)
                sistema.generar_reporte_ventas(ventas, f'reporte_ventas')
                archivo_adjunto = f'reporte_ventas.docx'
                destinatario = '' #Se elimino el correo destinatario por seguridad 
                asunto = 'Reporte de Ventas por Producto'
                mensaje = 'Adjunto se encuentra el reporte de ventas por producto.'
                sistema.enviar_correo_con_adjunto(destinatario, asunto, mensaje, archivo_adjunto)
                

        elif opcion == "5":
            print("Saliendo del sistema de ventas.")
            sistema.cerrar_conexion()
            break
        else:
            print("Opción no válida. Intente de nuevo.")


if __name__ == "__main__":
    main()


